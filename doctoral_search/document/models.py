import hashlib
from elasticsearch import Elasticsearch
from django.db import models
import io
from django.core.files.base import ContentFile
from PyPDF2 import PdfReader
from docx import Document
import magic 

ELASTIC_HOST = "http://localhost:9200"
INDEX_NAME = "research_library"
es_client = Elasticsearch(hosts=[ELASTIC_HOST])

class ResearchDocument(models.Model):
    title = models.CharField(max_length=255, verbose_name="Title", blank=False, null=False)
    description = models.TextField(verbose_name="Description", blank=True, null=True)
    author = models.CharField(max_length=255, verbose_name="Author", blank=False, null=False)
    keywords = models.TextField(help_text="put between each tow keywords ',' without any spaces")  # Store as comma-separated text
    #keywords_json = models.JSONField(default=list, blank=True, null=True) # Store as a list
    publication_date = models.DateField(blank=False, null=False)
    university = models.CharField(max_length=255, blank=False, null=False)

    file = models.FileField(upload_to='documents/')

    created_at = models.DateTimeField(auto_now_add=True)

    def get_keywords(self):
        """Return keywords as a list"""
        return self.keywords
    
    def get_absolute_url(self):
        return f"{self.file.path}" #reverse("model_detail", kwargs={"pk": self.pk})

    def compute_file_hash(self):
        """
        Safely computes MD5 hash of the file content
        Returns: MD5 hexdigest or None if error occurs
        """
        try:
            # Method 1: Using context manager (recommended)
            with self.file.open('rb') as f:
                file_hash = hashlib.md5()
                for chunk in f.chunks(4096):  # Read in chunks to handle large files
                    file_hash.update(chunk)
                return file_hash.hexdigest()
                
            # Method 2: Alternative one-line version (for small files)
            # with self.file.open('rb') as f:
            #     return hashlib.md5(f.read()).hexdigest()
                
        except (ValueError, IOError) as e:
            print(f"Error computing hash: {str(e)}")
            return None
    
    def extract_file_size(self):
        """Extract File Size"""
        return self.file.size

    def extract_file_type(self):    
        """Extract File Type"""
        return self.file.name.split('.')[-1].lower()
    
    def get_file_content(self):
        """
        Universal file content reader with automatic type detection
        Returns: 'Content'
        """
        result = ''

        try:
            # Detect file type using python-magic
            with self.file.open('rb') as f:
                file_bytes = f.read(2048)  # Read first 2KB for magic number detection
                mime = magic.from_buffer(file_bytes, mime=True)
                f.seek(0)  # Rewind for full read

                # Text files (including CSV, JSON, etc.)
                if mime.startswith('text/') or mime in [
                    'application/json',
                    'application/xml'
                ]:
                    try:
                        result = f.read().decode('utf-8')
                    except UnicodeDecodeError:
                        f.seek(0)
                        result = f.read().decode('latin-1')

                # PDF files
                elif mime == 'application/pdf':
                    pdf = PdfReader(f)
                    result= "\n".join(
                        page.extract_text() for page in pdf.pages
                        if page.extract_text()
                    )

                # Word documents
                elif mime in [
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'application/msword'
                ]:
                    doc = Document(f)
                    result= "\n".join(
                        para.text for para in doc.paragraphs
                    )

                # Binary files (images, executables, etc.)
                else:
                    result= file_bytes + f.read()  # Complete the read

        except Exception as e:
            result['error'] = str(e)

        return result

    def __str__(self):
        return self.title
